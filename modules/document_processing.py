
import os, re
from typing import Tuple
from docx import Document
from PyPDF2 import PdfReader

def extract_text_from_upload(file_path: str, file_type: str) -> Tuple[str, int]:
    file_type = file_type.lower().replace(".", "")
    if file_type == "txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read(), 1
    if file_type == "docx":
        doc = Document(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        return text, max(1, len(doc.paragraphs) // 25)
    if file_type == "pdf":
        reader = PdfReader(file_path)
        pages = []
        for p in reader.pages:
            pages.append(p.extract_text() or "")
        return "\n".join(pages), len(reader.pages)
    raise ValueError("Formato no permitido. Usa PDF, DOCX o TXT.")

def detect_sections(text: str):
    text_low = text.lower()
    candidates = [
        "carátula", "caratula", "índice", "indice", "resumen", "introducción", "introduccion",
        "planteamiento del problema", "objetivos", "objetivo general", "objetivos específicos",
        "justificación", "justificacion", "marco teórico", "marco teorico", "antecedentes",
        "hipótesis", "hipotesis", "metodología", "metodologia", "resultados",
        "discusión", "discusion", "conclusiones", "referencias", "bibliografía", "bibliografia"
    ]
    found = []
    for c in candidates:
        if c in text_low:
            found.append(c)
    return sorted(set(found))

def extract_references(text: str):
    refs = []
    marker = re.search(r"(referencias|bibliografía|bibliografia)(.*)$", text, flags=re.I|re.S)
    source = marker.group(2) if marker else text[-4000:]
    for line in source.splitlines():
        line = line.strip()
        if len(line) > 25 and re.search(r"\(\d{4}\)|\b(19|20)\d{2}\b", line):
            refs.append(line)
    # fallback: split by period author-like
    if not refs:
        chunks = re.split(r"\n|(?<=\.)\s+(?=[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+,\s)", source)
        refs = [c.strip() for c in chunks if len(c.strip()) > 45 and re.search(r"\b(19|20)\d{2}\b", c)]
    return refs[:50]

def estimate_academic_style(text: str):
    words = re.findall(r"\w+", text.lower())
    if not words:
        return 0
    academic_terms = ["investigación","metodología","objetivo","variable","hipótesis","población",
                      "muestra","instrumento","resultados","discusión","antecedente","marco","teórico",
                      "evidencia","análisis","enfoque","diseño"]
    count = sum(1 for w in words if w in academic_terms)
    return min(100, (count / max(len(words), 1)) * 1200)
