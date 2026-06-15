"""
modules/template_analyzer.py
Analiza plantillas DOCX/PDF para extraer su estructura y usarla como modelo.
"""
import re

# ── Detección de tipo de tabla ────────────────────────────────────────────────

_TABLE_TYPE_KEYWORDS = {
    "consistencia": ["problema", "objetivo", "hipotesis", "variable", "indicador",
                     "instrumento", "fuente", "metodologia", "diseño", "tipo"],
    "operacionalizacion": ["variable", "definicion", "dimension", "indicador",
                           "escala", "medicion", "items", "item", "instrumento"],
    "cronograma": ["actividad", "mes", "semana", "enero", "febrero", "marzo",
                   "abril", "mayo", "junio", "julio", "agosto", "setiembre",
                   "octubre", "noviembre", "diciembre", "periodo", "plazo"],
    "presupuesto": ["descripcion", "cantidad", "precio", "costo", "total",
                    "subtotal", "unidad", "monto", "inversion", "recurso"],
    "poblacion": ["institucion", "grado", "total", "poblacion", "muestra",
                  "seccion", "nivel", "numero", "estrato"],
    "estadisticos": ["media", "desviacion", "porcentaje", "frecuencia",
                     "n", "min", "max", "p-valor", "valor"],
}

# Secciones de portada que NO deben aparecer como contenido
_COVER_KEYWORDS = {
    'universidad', 'facultad', 'escuela', 'carrera', 'departamento',
    'proyecto de tesis', 'tesis para optar', 'trabajo de investigacion',
    'autor(es)', 'autores', 'asesor:', 'asesor :', 'linea de investigacion',
    'jurado dictaminador', 'jurado:', 'tribunal', 'ciudad', '— peru', '- peru',
    'titulo:', 'proyecto de investigacion cientifica', 'proyecto de tesis pregrado',
    'formato:', '(nombre de la', 'organización sincrónica', 'organización diacronica',
    'nombre de la institucion',
}

# Secciones que deben existir en la lista pero sin cuerpo de texto generado
_SKIP_CONTENT_KEYWORDS = {
    'indice general', 'tabla de contenido', 'contenido', 'indice de figuras',
    'lista de figuras', 'lista de tablas', 'indice de tablas',
    'referencias bibliograficas', 'referencias:', 'referencias',
    'anexos', 'anexo',  # el builder de anexos los maneja individualmente
}


def _normalize(text: str) -> str:
    text = text.lower()
    for src, dst in [('á','a'),('é','e'),('í','i'),('ó','o'),('ú','u'),('ñ','n')]:
        text = text.replace(src, dst)
    return text


def _parse_instructions(text: str) -> dict:
    """
    Extrae requisitos explícitos del texto instruccional de una sección de plantilla.
    Retorna un dict con los requisitos identificados.
    """
    n = _normalize(text)
    inst = {
        'raw':              text,
        'word_count':       None,
        'tense':            None,
        'citation_style':   None,
        'min_refs':         None,
        'max_refs':         None,
        'pct_scientific':   None,
        'years_recency':    None,
        'required_elements': [],
        'num_paragraphs':   None,
        'paragraph_lines':  None,
        'single_paragraph': False,
        'numbered_sections': False,
    }

    # Conteo de palabras: "150 a 200 palabras" / "100 palabras"
    m = re.search(r'(\d+)\s+a\s+(\d+)\s+palabras?', n)
    if m:
        inst['word_count'] = {'min': int(m.group(1)), 'max': int(m.group(2))}
    else:
        m = re.search(r'(\d+)\s+palabras?', n)
        if m:
            inst['word_count'] = {'min': int(m.group(1)), 'max': int(m.group(1))}

    # Tiempo verbal
    if 'tiempo pasado' in n:
        inst['tense'] = 'past'
    elif 'tiempo presente' in n:
        inst['tense'] = 'present'

    # Párrafo único
    if 'un solo parrafo' in n or 'un parrafo' in n:
        inst['single_paragraph'] = True

    # Número de párrafos / líneas
    m = re.search(r'no\s+superar\s+(\d+)\s+parrafo', n)
    if m:
        inst['num_paragraphs'] = int(m.group(1))
    m = re.search(r'menor\s+a\s+(\d+)\s+lineas?', n)
    if m:
        inst['paragraph_lines'] = int(m.group(1))

    # Estilo de citación
    for style in ['apa 7', 'apa7', 'apa 6', 'apa', 'ieee', 'vancouver', 'chicago', 'apa septima']:
        if style in n:
            inst['citation_style'] = style.upper().replace('APA7','APA 7').replace('APA SEPTIMA','APA 7')
            break

    # Mínimo de referencias
    m = re.search(r'(minimo|al menos|no menos de)\s+(\d+)\s+referencia', n)
    if m:
        inst['min_refs'] = int(m.group(2))

    # Porcentaje artículos científicos
    m = re.search(r'(\d+)\s*%\s+.*?articulo', n)
    if m:
        inst['pct_scientific'] = int(m.group(1))

    # Antigüedad de referencias (últimos N años)
    m = re.search(r'ultimos?\s+(\d+)\s+a[nñ]os?', n)
    if m:
        inst['years_recency'] = int(m.group(1))

    # Subtítulos numerados ("se enumerarán progresivamente")
    if 'numeraran' in n or 'numerados' in n or 'enumeraran' in n:
        inst['numbered_sections'] = True

    # Elementos requeridos: buscar listas de conceptos después de ":"
    # Patrón: "incluyendo: X, Y, Z" / "debe incluir: X, Y, Z"
    m = re.search(r'(?:incluyendo|debe incluir|incluir|contener)[:\s]+([^.]{10,200})', n)
    if m:
        raw_elements = re.split(r'[,;]', m.group(1))
        inst['required_elements'] = [e.strip() for e in raw_elements if len(e.strip()) > 3]

    return inst


def _is_cover_section(text: str) -> bool:
    n = _normalize(text.strip())
    return any(kw in n for kw in _COVER_KEYWORDS)


def _detect_table_type(headers: list, title: str = "") -> str:
    text = _normalize(" ".join(headers + [title]))
    scores = {}
    for ttype, keywords in _TABLE_TYPE_KEYWORDS.items():
        score = sum(1 for kw in keywords if kw in text)
        if score > 0:
            scores[ttype] = score
    if scores:
        return max(scores, key=scores.get)
    return "generic"


# ── Análisis DOCX ─────────────────────────────────────────────────────────────

def _analyze_docx(file_path: str) -> dict:
    from docx import Document
    doc = Document(file_path)

    sections = []
    tables_info = []
    last_real_section_idx = -1
    # Acumula el texto de los párrafos de contenido para cada sección
    content_buffers: list[list[str]] = []

    _CHAPTER_PAT = re.compile(
        r'^(cap[ií]tulo\s+[ivxIVX\d]+[\.:].{0,80}|'
        r'anexo\s+\d+[\.:].{0,80}|'
        r'\d+\.?\s{1,4}[A-ZÁÉÍÓÚÑ\w].{3,80}|'
        r'[IVX]+\.\s+[A-ZÁÉÍÓÚÑ].{3,60})$',
        re.IGNORECASE | re.UNICODE,
    )

    def _is_heading_style(style_name: str) -> tuple:
        sn = style_name.lower()
        for key in ('heading 1', 'título 1', 'titulo 1', 'heading1'):
            if key in sn: return True, 1
        for key in ('heading 2', 'título 2', 'titulo 2', 'heading2'):
            if key in sn: return True, 2
        for key in ('heading 3', 'título 3', 'titulo 3', 'heading3'):
            if key in sn: return True, 3
        return False, 0

    def _infer_level(text: str) -> int:
        n = _normalize(text.strip())
        if re.match(r'^(cap[ií]tulo|anexo)\s+', n):
            return 1
        if re.match(r'^\d+\s{2,}', text):
            return 1
        if re.match(r'^\d+\.\d+\s+', text):
            return 2
        if re.match(r'^\d+\.\d+\.\d+\s+', text):
            return 3
        if re.match(r'^[IVX]+\.\s+', text):
            return 1
        return 2

    def _add_section(level: int, title: str):
        nonlocal last_real_section_idx
        sections.append({"level": level, "title": title,
                         "content_preview": "", "full_content": "",
                         "instructions": {}})
        content_buffers.append([])
        last_real_section_idx = len(sections) - 1

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue

        is_heading, level = _is_heading_style(style_name)

        if is_heading:
            if not _is_cover_section(text):
                _add_section(level, text)
            continue

        if text.isupper() and 5 < len(text) < 140:
            if _is_cover_section(text):
                continue
            _add_section(_infer_level(text), text)
            continue

        if _CHAPTER_PAT.match(text) and len(text) < 140:
            if _is_cover_section(text):
                continue
            if not sections or _normalize(sections[-1]['title']) != _normalize(text):
                _add_section(_infer_level(text), text)
            continue

        is_bold_short = (
            len(text) < 100 and
            para.runs and
            all(r.bold for r in para.runs if r.text.strip())
        )
        if is_bold_short and not _is_cover_section(text):
            _add_section(3, text)
            continue

        # Párrafo de contenido — acumular en el buffer de la sección actual
        if last_real_section_idx >= 0:
            buf = content_buffers[last_real_section_idx]
            # Limitar a 4000 chars totales para no explotar el contexto de OpenAI
            total_so_far = sum(len(t) for t in buf)
            if total_so_far < 4000:
                buf.append(text)

    # Finalizar: volcar buffers en full_content + content_preview + instructions
    for i, sec in enumerate(sections):
        full = "\n".join(content_buffers[i]).strip()
        sec["full_content"]    = full[:4000]
        sec["content_preview"] = full[:200]
        sec["instructions"]    = _parse_instructions(full) if full else {}

    # Extraer requisitos globales del documento (ej. estilo de citación)
    all_content = "\n".join(s["full_content"] for s in sections)
    global_instructions = _parse_instructions(all_content)

    for table in doc.tables:
        rows_count = len(table.rows)
        headers = []
        if rows_count > 0:
            for cell in table.rows[0].cells:
                headers.append(cell.text.strip())
        table_type = _detect_table_type(headers)
        tables_info.append({
            "title": "",
            "headers": [h for h in headers if h],
            "rows": rows_count,
            "cols": len(table.columns),
            "type": table_type,
        })

    return {
        "sections": sections,
        "tables": tables_info,
        "has_cover": _has_cover(doc),
        "has_abstract": _has_section(sections, ["resumen", "abstract"]),
        "doc_type_hint": _detect_doc_type(sections, tables_info),
        "global_instructions": global_instructions,
    }


def _has_cover(doc) -> bool:
    """Detecta si el DOCX tiene portada revisando los primeros 15 párrafos."""
    cover_kw = {'universidad', 'facultad', 'escuela', 'proyecto de tesis',
                'trabajo de investigacion', 'tesis para optar'}
    count = 0
    for para in list(doc.paragraphs)[:15]:
        n = _normalize(para.text)
        if any(k in n for k in cover_kw):
            count += 1
    return count >= 2


# ── Análisis PDF ──────────────────────────────────────────────────────────────

def _analyze_pdf(file_path: str) -> dict:
    from PyPDF2 import PdfReader
    reader = PdfReader(file_path)

    all_text = ""
    for page in reader.pages:
        all_text += (page.extract_text() or "") + "\n"

    sections = _extract_headings_from_text(all_text)
    global_instructions = _parse_instructions(all_text[:8000])

    return {
        "sections": sections,
        "tables": [],
        "has_cover": False,
        "has_abstract": _has_section(sections, ["resumen", "abstract"]),
        "doc_type_hint": _detect_doc_type(sections, []),
        "global_instructions": global_instructions,
    }


def _extract_headings_from_text(text: str) -> list:
    lines = text.split('\n')
    sections = []
    heading_indices = []

    heading_patterns = [
        (1, r'^(CAP[IÍ]TULO\s+[IVX\d]+.{0,60})$'),
        (1, r'^(ANEXO\s+\d+.{0,60})$'),
        (1, r'^([IVX]+\.\s+[A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s]{5,50})$'),
        (1, r'^(\d+\s{2,}[A-ZÁÉÍÓÚÑ].{5,60})$'),
        (2, r'^(\d+\.\d+\s+.{5,80})$'),
        (3, r'^(\d+\.\d+\.\d+\s+.{5,80})$'),
    ]

    for i, line in enumerate(lines):
        line = line.strip()
        if not line or len(line) < 5 or _is_cover_section(line):
            continue
        for level, pattern in heading_patterns:
            if re.match(pattern, line, re.UNICODE):
                sections.append({"level": level, "title": line,
                                  "content_preview": "", "full_content": "",
                                  "instructions": {}})
                heading_indices.append(i)
                break
        else:
            if line.isupper() and 5 < len(line) < 120 and not _is_cover_section(line):
                sections.append({"level": 2, "title": line,
                                  "content_preview": "", "full_content": "",
                                  "instructions": {}})
                heading_indices.append(i)

    # Asignar contenido a cada sección (texto entre este heading y el siguiente)
    for j, (sec, hi) in enumerate(zip(sections, heading_indices)):
        next_hi = heading_indices[j + 1] if j + 1 < len(heading_indices) else len(lines)
        body_lines = [l.strip() for l in lines[hi+1:next_hi] if l.strip()]
        full = " ".join(body_lines)[:4000]
        sec["full_content"]    = full
        sec["content_preview"] = full[:200]
        sec["instructions"]    = _parse_instructions(full) if full else {}

    return sections


# ── Helpers ───────────────────────────────────────────────────────────────────

def _has_section(sections: list, keywords: list) -> bool:
    for s in sections:
        norm = _normalize(s["title"])
        if any(kw in norm for kw in keywords):
            return True
    return False


def _detect_doc_type(sections: list, tables: list) -> str:
    all_titles = " ".join(_normalize(s["title"]) for s in sections)
    table_types = {t.get("type", "") for t in tables}

    # Señales fuertes de tesis/proyecto (capítulos numerados, anexos)
    has_chapters = any(kw in all_titles for kw in ["capitulo i", "capitulo ii", "capitulo iii",
                                                     "capitulo iv", "capitulo v"])
    has_annexes  = "anexo" in all_titles

    # Señales fuertes de artículo (sin capítulos, con secciones propias de artículo)
    article_exclusive = any(kw in all_titles for kw in [
        "materiales y metodo", "resultados y discusion", "resultados de la revision",
        "conflicto de interes", "contribucion de autoria",
    ])

    if article_exclusive and not has_chapters:
        return "articulo"

    if has_chapters or has_annexes:
        # Distinguir tesis (5 caps) de proyecto (3 caps + admin)
        if any(kw in all_titles for kw in ["aspectos administrativos", "cronograma",
                                            "presupuesto", "planteamiento"]):
            return "proyecto_tesis"
        if "presupuesto" in table_types or "cronograma" in table_types:
            return "proyecto_tesis"
        if any(kw in all_titles for kw in ["resultado", "discusion", "hallazgo", "capitulo iv"]):
            return "tesis"
        return "proyecto_tesis"

    # Sin capítulos y sin señales de artículo → inferir por tablas
    if "presupuesto" in table_types or "cronograma" in table_types:
        return "proyecto_tesis"

    # Abstract/keywords solos (sin capítulos) → artículo
    if any(kw in all_titles for kw in ["abstract", "keywords", "palabras clave"]):
        return "articulo"

    return "proyecto_tesis"


def _empty_structure() -> dict:
    return {
        "sections": [],
        "tables": [],
        "has_cover": False,
        "has_abstract": False,
        "doc_type_hint": "proyecto_tesis",
        "global_instructions": {},
    }


# ── Punto de entrada público ──────────────────────────────────────────────────

def analyze_template(file_path: str, file_type: str) -> dict:
    """
    Analiza una plantilla DOCX o PDF y devuelve su estructura.

    Returns dict con:
        sections      — lista de {level, title, content_preview}
        tables        — lista de {title, headers, rows, cols, type}
        has_cover     — bool
        has_abstract  — bool
        doc_type_hint — "tesis" | "proyecto_tesis" | "articulo"
    """
    file_type = file_type.lower().replace(".", "")
    try:
        if file_type == "docx":
            return _analyze_docx(file_path)
        elif file_type == "pdf":
            return _analyze_pdf(file_path)
    except Exception as e:
        print(f"[template_analyzer] Error analizando plantilla: {e}")
    return _empty_structure()
